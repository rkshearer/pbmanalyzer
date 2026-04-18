import io
import pdfplumber
from docx import Document as DocxDocument


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                raise ValueError("The PDF file contains no pages.")
            for i, page in enumerate(pdf.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(f"--- Page {i + 1} ---\n{text.strip()}")
                except Exception:
                    continue
        return "\n\n".join(text_parts)
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(
            f"Could not read the PDF file. It may be corrupted or password-protected. ({type(e).__name__})"
        )


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(c for c in row_cells if c)
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(
            f"Could not read the DOCX file. It may be corrupted or in an unsupported format. ({type(e).__name__})"
        )


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Please upload a PDF or DOCX file.")
