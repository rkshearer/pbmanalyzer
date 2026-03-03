"""
Negotiation letter generator.

Calls Claude to draft a professional PBM contract negotiation letter based on
the contract analysis, then formats it as a DOCX document using python-docx.
"""

import io
import os
from datetime import date

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .models import PBMAnalysisReport

_LETTER_PROMPT = """\
You are drafting a professional PBM contract negotiation letter on behalf of an employee benefits consultant.

Based on the contract analysis below, write a formal negotiation letter addressed to the PBM.

REQUIREMENTS:
1. Open with today's date, a recipient address block, and a clear subject line
2. State that this is a formal request for contract modifications following an independent analysis
3. For each underperforming pricing term, write a clear paragraph that:
   - Names the specific term
   - States what the contract currently provides
   - States the market benchmark
   - Makes a specific, numeric improvement request (e.g., "We request an improvement to AWP-82% from the current AWP-74%")
4. Address structural concerns (MAC transparency, audit rights, rebate pass-through, etc.) with specific asks
5. Set a firm 30-day response deadline from today ({today})
6. Close professionally with a signature block

PLACEHOLDERS TO USE:
- [BROKER NAME] — the consultant's full name
- [BROKER FIRM] — the consulting firm name
- [CLIENT NAME] — the employer plan sponsor
- [PBM CONTACT NAME] — the PBM account representative

CONTRACT ANALYSIS:
Grade: {grade}
Parties: {parties}
Contract Term: {term}

PRICING vs. MARKET:
- Brand Retail AWP: {brand_retail} (benchmark: {brand_benchmark}, assessment: {brand_assessment})
- Generic Retail AWP: {generic_retail} (benchmark: {generic_benchmark}, assessment: {generic_assessment})
- Specialty AWP: {specialty} (benchmark: {specialty_benchmark}, assessment: {specialty_assessment})
- Retail Dispensing Fee: {dispensing_fee}
- Rebate Guarantee: {rebate}
- MAC Terms: {mac}
- Admin Fees: {admin_fees}

KEY CONCERNS:
{concerns}

NEGOTIATION GUIDANCE FROM ANALYSIS:
{guidance}

Write the complete letter text only — no preamble or explanation outside the letter itself.
"""


def generate_negotiation_letter(analysis: PBMAnalysisReport) -> bytes:
    """Generate a negotiation letter DOCX from the contract analysis. Returns raw bytes."""
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    mc = analysis.market_comparison
    pt = analysis.pricing_terms

    prompt = _LETTER_PROMPT.format(
        today=date.today().strftime("%B %d, %Y"),
        grade=analysis.overall_grade,
        parties=analysis.contract_overview.parties,
        term=analysis.contract_overview.contract_term,
        brand_retail=pt.brand_retail_awp_discount,
        brand_benchmark=mc.brand_retail_benchmark,
        brand_assessment=mc.brand_retail_assessment,
        generic_retail=pt.generic_retail_awp_discount,
        generic_benchmark=mc.generic_retail_benchmark,
        generic_assessment=mc.generic_retail_assessment,
        specialty=pt.specialty_awp_discount,
        specialty_benchmark=mc.specialty_benchmark,
        specialty_assessment=mc.specialty_assessment,
        dispensing_fee=pt.retail_dispensing_fee,
        rebate=pt.rebate_guarantee,
        mac=pt.mac_pricing_terms,
        admin_fees=pt.admin_fees,
        concerns="\n".join(f"- {c}" for c in analysis.key_concerns),
        guidance="\n".join(f"{i + 1}. {g}" for i, g in enumerate(analysis.negotiation_guidance)),
    )

    response = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}],
    )

    letter_text = response.content[0].text
    return _build_docx(letter_text, analysis)


def _build_docx(letter_text: str, analysis: PBMAnalysisReport) -> bytes:
    """Format the letter text into a professional DOCX document."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Document header
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("PBM CONTRACT NEGOTIATION LETTER")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f"Contract Grade: {analysis.overall_grade}  ·  "
        f"Generated: {date.today().strftime('%B %d, %Y')}  ·  Confidential"
    )
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    div = doc.add_paragraph("─" * 72)
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.runs[0].font.color.rgb = RGBColor(0xC9, 0x96, 0x0F)

    doc.add_paragraph()

    # Letter body — preserve paragraph structure
    for line in letter_text.split("\n"):
        para = doc.add_paragraph()
        if line.strip():
            run = para.add_run(line)
            run.font.size = Pt(11)

    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(
        "Generated by PBM Contract Analyzer  ·  Confidential  ·  Not Legal Advice"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
